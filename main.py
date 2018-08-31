#!/usr/bin/env python3

import sys
import argparse
import logging
from datetime import datetime
import serial
import csv
from docx import Document
from docx.shared import Cm
from docx.enum.table import WD_TABLE_ALIGNMENT


def parse_data(data: str):
    parsed = {}
    converted = {}
    skip = ['frame-head', 'number-mark', 'WBC-warning', 'reserved',
            'WBC-hist-scale-line-pos', 'RBC-hist-scale-line-pos',
            'PLT-hist-scale-line-pos', 'reserved2', 'WBC-y-data',
            'RBC-y-data', 'PLT-y-data', 'frame-end', 'RBC-dec', 'omission',
            'RcdNo']
    for key in char_positions.keys():
        if key not in skip:
            start, end = char_positions[key]
            parsed[key] = data[start:end]
    for key in data_parsers.keys():
        if key not in skip:
            converted[key] = data_parsers[key](parsed[key])
    return converted


char_positions = {
    'id': (2096, 2105),
    'name': (2106, 2120),
    'frame-head': (0, 2),
    'number-mark': (2, 11),
    'test-date': (11, 25),
    'WBC': (25, 29),
    'LYM%': (29, 33),
    'MID%': (33, 37),
    'NEUT%': (37, 41),
    'LYM#': (41, 45),
    'MID#': (45, 49),
    'NEUT#': (49, 53),
    'RBC': (53, 57),
    'HGB': (57, 61),
    'HCT': (61, 65),
    'MCV': (65, 70),
    'MCH': (70, 75),
    'MCHC': (75, 80),
    'RDW-SD': (80, 84),
    'RDW-CV': (84, 88),
    'PLT': (88, 92),
    'MPV': (92, 96),
    'PDW': (96, 100),
    'PCT': (100, 104),
    'P-LCR': (104, 108),
    'RBC-dec': (108, 109),
    'omission': (109, 114),
    'RcdNo': (114, 129),
    'WBC-warning': (129, 134),
    'reserved': (134, 143),
    'WBC-hist-scale-line-pos': (143, 155),
    'RBC-hist-scale-line-pos': (155, 161),
    'PLT-hist-scale-line-pos': (161, 167),
    'reserved2': (167, 175),
    'WBC-y-data': (175, 943),
    'RBC-y-data': (943, 1711),
    'PLT-y-data': (1711, 2095),
    'frame-end': (2095, 2096)
}

data_parsers = {
    'id': lambda x: [x],
    'name': lambda x: [x.split('#')[0]],
    'frame-head': lambda x: [x] if (x == '@a') else ValueError,
    'number-mark': lambda x: [x],
    'test-date': lambda x: [str(datetime.strptime(x, '%Y%m%d%H%M%S'))],
    'WBC': lambda x: [x[0:2] + '.' + x[2] + ' ' + x[3]],
    'LYM%': lambda x: [x[0:2] + '.' + x[2] + '%'],
    'MID%': lambda x: [x[0:2] + '.' + x[2] + '%'],
    'NEUT%': lambda x: [x[0:2] + '.' + x[2] + ' ' + x[3]],
    'LYM#': lambda x: [x[0:2] + '.' + x[2]],
    'MID#': lambda x: [x[0:2] + '.' + x[2]],
    'NEUT#': lambda x: [x[0:2] + '.' + x[2] + ' ' + x[3]],
    'RBC': lambda x: [x[0:1] + '.' + x[1:3] + ' ' + x[3]],
    'HGB': lambda x: [x[0:3] + ' ' + x[3]],
    'HCT': lambda x: [x[0:2] + '.' + x[2] + ' ' + x[3]],
    'MCV': lambda x: [x[0:3] + '.' + x[3] + ' ' + x[4]],
    'MCH': lambda x: [x[0:3] + '.' + x[3]],
    'MCHC': lambda x: [x[0:4] + ' ' + x[4]],
    'RDW-SD': lambda x: [x[0:2] + '.' + x[2] + ' ' + x[3]],
    'RDW-CV': lambda x: [x[0:2] + '.' + x[2] + ' ' + x[3]],
    'PLT': lambda x: [x[0:3]],
    'MPV': lambda x: [x[0:2] + '.' + x[2:3]],
    'PDW': lambda x: [x[0:2] + '.' + x[2:3]],
    'PCT': lambda x: [x[0] + '.' + x[1:3]],
    'P-LCR': lambda x: [x[0:2] + '.' + x[2:3]],
    'RBC-dec': lambda x: [x],
    'omission': lambda x: [x],
    'RcdNo': lambda x: [x],
    'WBC-warning': lambda x: [x],
    'reserved': lambda x: [x],
    'WBC-hist-scale-line-pos': lambda x: [int(x[0:3]), int(x[3:6]), int(x[6:9]), int(x[9:])],
    'RBC-hist-scale-line-pos': lambda x: [int(x[0:3]), int(x[3:6])],
    'PLT-hist-scale-line-pos': lambda x: [int(x[0:3]), int(x[3:6])],
    'reserved2': lambda x: [x],
    'WBC-y-data': lambda x: [int(x[start:start+3]) for start in range(0, len(x), 3)],
    'RBC-y-data': lambda x: [int(x[start:start+3]) for start in range(0, len(x), 3)],
    'PLT-y-data': lambda x: [int(x[start:start+3]) for start in range(0, len(x), 3)],
    'frame-end': lambda x: [x] if (x == '#') else ValueError
}


def transpose(cols):
    def mypop(l):
        try:
            return l.pop(0)
        except IndexError:
            return ''
    while any(cols):
        yield [mypop(l) for l in cols]


def generate_reference():
    return [
        'Reference',
        '',
        '',
        '',
        '10*9/L    6.0 -- 17.0',
        '%        12.0 -- 30.0',
        '%         5.0 -- 20.0',
        '%        60.0 -- 70.0',
        '10*9/L    1.0 --  4.8',
        '10*9/L    0.2 --  2.1',
        '10*9/L    3.0 -- 11.4',
        '10*12/L   5.5 --  8.5',
        'g/dL     12.0 -- 18.0',
        '%        37.0 -- 55.0',
        'fL       60.0 -- 70.0',
        'pg       19.5 -- 24.5',
        'g/dL     32.0 -- 36.0',
        'fL       37.0 -- 54.0',
        '%        11.0 -- 15.5',
        '10*9/L    200 --  900',
        'fL        7.0 -- 12.0',
        '%         9.0 -- 30.0',
        '%         0.1 --  9.9',
        '%         9.0 -- 50.0',
    ]


def main(args, loglevel):
    logging.basicConfig(format="%(levelname)s: %(message)s", level=loglevel)

    logging.info(f'Indicated COM PORT is {args.com_port}')
    ser = serial.Serial(args.com_port, 9600)
    while True:
        logging.info('Ready to receive data! Waiting..')
        data = ser.read(2120)

        logging.info('Data has been received!')
        data = "".join(map(chr, data))

        logging.info('Checking data integrity..')
        if not (data[0:2] == '@a' and data[-2] == '#'):
            logging.error(
                f'Please re-run! Got header: {data[0:2]},tail: {data[-2]}')
            raise ValueError

        logging.info('Parsing and converting data..')
        converted_data = parse_data(data)

        headers = list(converted_data.keys())
        headers = ['Composition'] + headers
        items = ['Measured'] + \
            list(map(lambda x: x[0], converted_data.values()))
        data_to_write = list(zip(headers, items, generate_reference()))
        now = str(datetime.now()).replace(':', '-')

        if args.is_docx:
            logging.info('Writing to docx file..')
            document = Document()

            sections = document.sections
            for section in sections:
                section.left_margin = Cm(1)
                section.right_margin = Cm(1)

            table = document.add_table(rows=0, cols=3)
            table.style = 'Light Grid Accent 1'
            table.alignment = WD_TABLE_ALIGNMENT.LEFT
            for label, value, ref in data_to_write:
                row_cells = table.add_row().cells
                row_cells[0].text = label
                row_cells[1].text = value
                row_cells[2].text = ref

            for row in table.rows:
                for cell in row.cells:
                    paragraphs = cell.paragraphs
                    for paragraph in paragraphs:
                        for run in paragraph.runs:
                            font = run.font
                            font.name = 'Courier'

            document.save(f'{converted_data["id"][0]}_{now}.docx')

        else:
            logging.info('Writing to CSV file..')
            with open(converted_data['id'][0] + '_' + now + '.csv', 'w', newline='') as out_file:
                writer = csv.writer(out_file, dialect='excel')
                writer.writerows(data_to_write)
            logging.info('Success!')


if __name__ == '__main__':
    parser = argparse.ArgumentParser(
        description="KT series communication protocol Reader")
    # TODO Specify your real parameters here.
    parser.add_argument(
        "com_port",
        help="indicate COM port",
        metavar="com_port")
    parser.add_argument(
        "-v",
        "--verbose",
        help="increase output verbosity",
        action="store_true")
    parser.add_argument(
        "--docx",
        "--doc",
        dest='is_docx',
        action='store_true',
        help="Generate to docx. This is the default",)
    parser.add_argument(
        "--csv",
        dest='is_docx',
        action='store_false',
        help="Generate to csv.",)
    parser.set_defaults(is_docx=True)
    args = parser.parse_args()

    # Setup logging
    if args.verbose:
        loglevel = logging.DEBUG
    else:
        loglevel = logging.INFO

    main(args, loglevel)
